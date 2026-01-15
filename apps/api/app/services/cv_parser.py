import pdfplumber
from docx import Document
from typing import Optional


class CVParser:
    """CV parser service (PDF та DOCX)"""
    
    @staticmethod
    async def parse_pdf(file_path: str) -> Optional[str]:
        """ PDF parser"""
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip() if text else None
        except Exception as e:
            raise Exception(f"Error parse PDF: {str(e)}")
    
    @staticmethod
    async def parse_docx(file_path: str) -> Optional[str]:
        """ DOCX parser """
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            text = "\n".join(text_parts)
            return text.strip() if text else None
        except Exception as e:
            raise Exception(f"Error parse DOCX: {str(e)}")
    
    @staticmethod
    async def parse_file(file_path: str, file_type: str) -> Optional[str]:
        """ Universal method for parsing files """
        file_type_lower = file_type.lower()
        
        if file_type_lower == "pdf":
            return await CVParser.parse_pdf(file_path)
        elif file_type_lower in ["docx", "doc"]:
            return await CVParser.parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}. Supported file types: pdf, docx, doc")
    
    @staticmethod
    def validate_file_type(file_type: str) -> bool:
        """ Check if file type is supported """
        supported_types = ["pdf", "docx", "doc"]
        return file_type.lower() in supported_types